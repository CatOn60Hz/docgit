import os
from docx import Document

# Create a sample document
doc = Document()
doc.add_heading('Hello World', 0)
doc.add_paragraph('This is a test document.')
doc.save('test.docx')

print("Created test.docx")
