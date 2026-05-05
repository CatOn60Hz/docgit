from docx import Document

doc = Document('test.docx')
doc.paragraphs[0].text = 'Hello World Modified!'
doc.add_paragraph('This is a new line added in the second version.')
doc.save('test.docx')

print("Modified test.docx")
