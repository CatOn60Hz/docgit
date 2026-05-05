from docx import Document
from docx.oxml.ns import qn

def extract_content(filepath):
    doc = Document(filepath)
    page_num = 1
    para_num = 1
    
    for para in doc.paragraphs:
        # Check for page breaks in the paragraph's XML
        # lastRenderedPageBreak or explicit page breaks
        for run in para._element.xpath('.//w:lastRenderedPageBreak | .//w:br[@w:type="page"]'):
            page_num += 1
            
        # Check for images
        images = para._element.xpath('.//w:drawing')
        image_tags = f" [Contains {len(images)} Image(s)]" if images else ""
        
        print(f"Page {page_num}, Line {para_num}: {para.text[:50]}{image_tags}")
        para_num += 1

# Create a test doc
doc = Document()
doc.add_paragraph("First paragraph")
doc.add_page_break()
doc.add_paragraph("Second paragraph on page 2")
doc.save('test_parse.docx')

extract_content('test_parse.docx')
